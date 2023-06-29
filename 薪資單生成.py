import tkinter as tk
from tkinter import filedialog, messagebox
import openpyxl as opxl
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
import os
import io
import sys
import msoffcrypto as mso
import win32com.client as win32


# 函數：在GUI介面中顯示文字
def show_text_to_GUI(text):
    lable_detail.insert(tk.END, text)
    window.update()

# 函數：加載加密的Excel檔案
def load_encrypted_excel(encrypted_filename: str):
    decrypted = io.BytesIO()
    password = input_password.get()
    with open(encrypted_filename, "rb") as f:
        file = mso.OfficeFile(f)
        file.load_key(password=password)
        file.decrypt(decrypted)
    return decrypted

# 函數：替換Word檔案中的文字
def replace_word_text(word_file, old_text: str, new_text: str):
    try:
        for paragraph in word_file.paragraphs:
            paragraph.text = paragraph.text.replace(old_text, new_text)

        for table in word_file.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = cell.text.replace(old_text, new_text)
    except:
        show_text_to_GUI('填入錯誤\n')

# 函數：從密碼表中獲取密碼
def get_password(path_password):
    try:
        show_text_to_GUI('正在取得密碼\n')
        password_workbook = opxl.load_workbook(path_password)
        password_sheet = password_workbook.worksheets[0]
        password_dic = {
            password_sheet.cell(row=pwrow, column=1).value: password_sheet.cell(row=pwrow, column=2).value
            for pwrow in range(1, password_sheet.max_row + 1)}
        show_text_to_GUI('取得密碼成功\n')
        return password_dic
    except:
        show_text_to_GUI('取得密碼失敗\n')
        raise

# 函數：加密檔案
def encrypt_files(file_path, password_dic: dict, name):
    try:
        doc = word.Documents.Open(file_path)
        doc.Password = str(password_dic.get(name))
        doc.Save()
        doc.Close()
        show_text_to_GUI(f'{name}薪資單已加密完畢\n')
    except KeyError:
        pass
    except:
        show_text_to_GUI('加密失敗\n')
        raise

# 函數：選擇檔案路徑
def files_path():
    global source_file, path_docx, path_password, generate_path
    try:
        root = tk.Tk()
        root.withdraw()

        source_file = filedialog.askopenfilename(title='選擇轉換檔案')
        if not source_file:  # 如果沒有選擇檔案，直接返回
            return
        path_docx = filedialog.askopenfilename(title='選擇範例word')
        if not path_docx:  # 如果沒有選擇檔案，直接返回
            return
        path_password = filedialog.askopenfilename(title='選擇密碼表')
        if not path_password:  # 如果沒有選擇檔案，直接返回
            return
        generate_path = filedialog.askdirectory(title='選擇生成檔案資料夾')
        if not generate_path:  # 如果沒有選擇檔案，直接返回
            return

    except:
        show_text_to_GUI('檢查資料夾中檔案是否正確\n')
        messagebox.showerror("錯誤", "檢查資料夾中的檔案是否正確")
        raise

# 函數：加載檔案
def load_files():
    global sheet, paydate
    try:
        workbook = opxl.load_workbook(source_file, data_only=True)
        sheet = workbook['薪資總表']
        paydate = str(sheet['A2'].value)[:11]
    except:
        try:
            decryptedwb = load_encrypted_excel(source_file)
            workbook = opxl.load_workbook(decryptedwb, data_only=True)
            sheet = workbook['薪資總表']
            paydate = str(sheet['A2'].value)[:11]
        except:
            show_text_to_GUI('讀取錯誤\n檢查EXCEL檔案或密碼\n')
            messagebox.showerror("錯誤", "讀取錯誤，請檢查EXCEL檔案或密碼")
            raise

# 函數：獲取使用者輸入
def get_user_input():
    global start_row, end_row
    while True:
        try:
            start_row = int(input_start_row.get())
            end_row = int(input_end_row.get())
            messagebox.showinfo("提示", "關閉 Excel 檔案後按下 Enter 鍵開始生成")
            break
        except ValueError:
            messagebox.showerror("錯誤", "請輸入整數！")

# 函數：替換表格中的內容
def table_replace():
    global word
    load_files()
    get_user_input()
    try:
        password = get_password(path_password)
        word = win32.Dispatch('Word.Application')
        for table_row in range(start_row + 1, end_row + 1):
            word_file = Document(path_docx)
            for table_col in range(3, sheet.max_column + 1):
                old_text = str(sheet.cell(row=start_row, column=table_col).value)
                new_value = sheet.cell(row=table_row, column=table_col).value
                new_text = str(new_value) if new_value not in [None, 0] else ""
                replace_word_text(word_file, old_text, new_text)
            name = str(sheet.cell(row=table_row, column=6).value)
            generate_file_name = f"{paydate}薪資單-{name}.docx"
            title_text = word_file.paragraphs[0].text.replace('#發薪日期', paydate)
            word_file.paragraphs[0].text = title_text
            title_text_style = word_file.paragraphs[0].runs[0]
            title_text_style.font.name = "Arial"
            title_text_style.bold = True
            title_text_style._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
            title_text_style.font.size = Pt(14)
            word_file.save(os.path.join(generate_path, generate_file_name))
            show_text_to_GUI(f'{name}薪資單生成完畢\n')
            path_unprotectdocx = os.path.normpath(os.path.join(generate_path, generate_file_name))
            encrypt_files(path_unprotectdocx, password, name)
        word.Quit()
        show_text_to_GUI('薪資單均已加密完畢\n')
        messagebox.showerror("完成", "完成!!")
    except:
        show_text_to_GUI('填入過程中出現錯誤\n')
        messagebox.showerror("錯誤", "填入過程中出現錯誤")
        raise

# 函數：關閉視窗
def close_window():
    window.destroy()
    sys.exit()

# 主函數
def main():
    global input_start_row, input_end_row, detail, window, lable_detail, input_password

    window = tk.Tk()
    window.title("薪資單生成程式")
    width = 600
    height = 400
    screen_width = window.winfo_screenwidth()
    screen_height = window.winfo_screenheight()
    x = (screen_width / 2) - (width / 2)
    y = (screen_height / 2) - (height / 2)
    window.geometry("%dx%d+%d+%d" % (width, height, x, y))
    window.resizable(0, 0)

    frame_input = tk.Frame(window)
    frame_input.pack(side=tk.TOP, pady=20)

    label_start_row = tk.Label(frame_input, text="起始列：")
    label_start_row.grid(row=0, column=0, sticky=tk.W, padx=10)
    input_start_row = tk.Entry(frame_input)
    input_start_row.grid(row=0, column=1, padx=10)

    label_end_row = tk.Label(frame_input, text="結束列：")
    label_end_row.grid(row=1, column=0, sticky=tk.W, padx=10)
    input_end_row = tk.Entry(frame_input)
    input_end_row.grid(row=1, column=1, padx=10)

    frame_password = tk.Frame(window)
    frame_password.pack(side=tk.TOP, pady=20)

    label_password = tk.Label(frame_password, text="加密密碼：")
    label_password.grid(row=0, column=0, sticky=tk.W, padx=10)
    input_password = tk.Entry(frame_password, show="*")
    input_password.grid(row=0, column=1, padx=10)

    frame_buttons = tk.Frame(window)
    frame_buttons.pack(side=tk.TOP, pady=20)

    button_select = tk.Button(frame_buttons, text="選擇檔案", command=files_path)
    button_select.grid(row=0, column=0, padx=10)

    button_generate = tk.Button(frame_buttons, text="生成薪資單", command=table_replace)
    button_generate.grid(row=0, column=1, padx=10)

    close_button = tk.Button(window, text="Close", command=close_window)
    close_button.pack()

    frame_detail = tk.Frame(window)
    frame_detail.pack(side=tk.TOP, pady=20)

    lable_detail = tk.Text(frame_detail, width=80, height=10)
    lable_detail.pack(side=tk.LEFT, fill=tk.Y)
    scrollbar = tk.Scrollbar(frame_detail, command=lable_detail.yview)
    scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
    lable_detail.config(yscrollcommand=scrollbar.set)

    window.protocol("WM_DELETE_WINDOW", close_window)
    window.mainloop()

if __name__ == '__main__':
    main()
